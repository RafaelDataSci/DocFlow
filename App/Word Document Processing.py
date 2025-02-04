#!/usr/bin/env python
"""
Enhanced Word Document Processing Script

- Processes DOCX files in parallel.
- For each file, replaces user-specified text pairs using deep run-splitting.
- Renames the file by applying the same replacements to the filename.
- Prompts the user for origin and destination folders.
- Displays progress messages so the user can follow along.

Note:
    New text replacements now explicitly copy the original text’s style (font type, size, etc.)
    to ensure the formatting is preserved.
"""

import logging
from pathlib import Path
from docx import Document
from concurrent.futures import ProcessPoolExecutor, as_completed

# Configure logging to show timestamps and log levels.
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)

def deep_replace_text_in_paragraph(paragraph, old_text, new_text):
    """
    Replace occurrences of old_text with new_text in a paragraph.
    
    First, attempts to replace matches within a single run by splitting that run.
    If old_text isn’t fully contained in any one run (i.e. it spans multiple runs),
    then the entire paragraph's text is replaced as a fallback.
    
    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph to process.
        old_text (str): The text to be replaced.
        new_text (str): The text to insert.
    """
    replaced_in_run = False
    for run in paragraph.runs:
        if old_text in run.text:
            parts = []
            start = 0
            while True:
                idx = run.text.find(old_text, start)
                if idx == -1:
                    parts.append(run.text[start:])
                    break
                parts.append(run.text[start:idx])
                parts.append(new_text)
                start = idx + len(old_text)
            # Replace the current run's text with the first segment.
            run.text = parts[0]
            # Get the parent element to allow insertion of new run elements.
            parent = run._element.getparent()
            run_index = list(parent).index(run._element)
            # Insert new runs for the remaining segments while copying formatting.
            for seg in parts[1:]:
                new_run = paragraph.add_run(seg)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.style = run.style  # Copy the style (ensures same font type, color, etc.)
                run_index += 1
                parent.insert(run_index, new_run._element)
            replaced_in_run = True
    # Fallback: if no single run contained the entire old_text, replace full paragraph text.
    if not replaced_in_run and old_text in paragraph.text:
        full_text = paragraph.text
        new_full_text = full_text.replace(old_text, new_text)
        p_element = paragraph._element
        for run in list(paragraph.runs):
            p_element.remove(run._element)
        paragraph.add_run(new_full_text)

def process_word_file(file_path: Path, replacements: list, destination_folder: Path) -> None:
    """
    Process a single DOCX file:
      - Open the document.
      - For each paragraph, apply all replacement pairs.
      - Save the document under a new filename (with replacements applied to the filename)
        to the destination folder.
    
    Args:
        file_path (Path): The DOCX file to process.
        replacements (list): List of (old_text, new_text) tuples.
        destination_folder (Path): Folder where the new DOCX file will be saved.
    """
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements:
                if old_text in paragraph.text:
                    deep_replace_text_in_paragraph(paragraph, old_text, new_text)
        # Update the filename by applying each replacement.
        new_filename = file_path.name
        for old_text, new_text in replacements:
            new_filename = new_filename.replace(old_text, new_text)
        new_file_path = destination_folder / new_filename
        doc.save(new_file_path)
        logging.info(f"Created new file: {new_file_path}")
    except Exception as e:
        logging.error(f"Error processing {file_path}: {e}")

def get_replacements_from_user() -> list:
    """
    Prompt the user to input replacement pairs.
    
    Returns:
        A list of tuples: (old_text, new_text). Input ends when the old_text is empty.
    """
    replacements = []
    print("Enter replacement pairs. (Leave the old text empty to finish.)")
    while True:
        old_text = input("Enter text to replace: ").strip()
        if not old_text:
            break
        new_text = input("Enter replacement text: ").strip()
        replacements.append((old_text, new_text))
    return replacements

def process_directory(origin_folder: Path, destination_folder: Path, replacements: list) -> None:
    """
    Process all DOCX files in the origin folder concurrently, saving output files to the destination folder,
    and display progress.
    
    Args:
        origin_folder (Path): Folder containing DOCX files.
        destination_folder (Path): Folder where processed DOCX files will be saved.
        replacements (list): List of (old_text, new_text) tuples.
    """
    files = [f for f in origin_folder.glob("*.docx") if f.is_file()]
    total_files = len(files)
    if total_files == 0:
        logging.info("No DOCX files found in the origin folder.")
        return
    logging.info(f"Found {total_files} DOCX files to process.")
    destination_folder.mkdir(parents=True, exist_ok=True)
    files_processed = 0
    with ProcessPoolExecutor() as executor:
        future_to_file = {
            executor.submit(process_word_file, file, replacements, destination_folder): file
            for file in files
        }
        completed = 0
        for future in as_completed(future_to_file):
            completed += 1
            file = future_to_file[future]
            try:
                future.result()
                files_processed += 1
                logging.info(f"Processed {completed}/{total_files}: {file.name}")
            except Exception as e:
                logging.error(f"Error processing {file.name}: {e}")
    logging.info(f"Processing complete! Total files processed: {files_processed}/{total_files}")

if __name__ == "__main__":
    origin_input = input("Enter the ORIGIN folder path containing DOCX files: ").strip()
    destination_input = input("Enter the DESTINATION folder path for updated DOCX files: ").strip()
    origin_folder = Path(origin_input)
    destination_folder = Path(destination_input)
    if not origin_folder.is_dir():
        logging.error(f"Origin folder does not exist: {origin_folder}")
    else:
        replacements = get_replacements_from_user()
        if not replacements:
            print("No replacements provided. Exiting.")
        else:
            process_directory(origin_folder, destination_folder, replacements)
