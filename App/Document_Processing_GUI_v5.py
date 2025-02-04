#!/usr/bin/env python
import subprocess
import sys
import threading
import logging
from pathlib import Path
from queue import Queue
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from PyPDF2 import PdfMerger
import PySimpleGUI as sg

# List of required dependencies
REQUIRED_LIBS = [
    "PySimpleGUI",
    "python-docx",
    "docx2pdf",
    "PyPDF2"
]

def check_and_install_dependencies():
    """Check and install missing dependencies automatically."""
    for lib in REQUIRED_LIBS:
        try:
            __import__(lib)
        except ImportError:
            print(f"⚠ {lib} not found. Installing...")
            try:
                subprocess.check_call([sys.executable, "-m", "pip", "install", lib])
                print(f"✅ {lib} installed successfully!")
            except Exception as e:
                print(f"❌ Failed to install {lib}. Please install it manually using:")
                print(f"   pip install {lib}")

# Run the dependency check
check_and_install_dependencies()

# GUI Setup
sg.theme('DarkBlue14')

# Layout for Word Processing
tab1_layout = [
    [sg.Text("Word Document Processing", font=("Any", 16))],
    [sg.Text("Origin Folder:"), sg.Input(key="-WP_ORIGIN-"), sg.FolderBrowse()],
    [sg.Text("Destination Folder:"), sg.Input(key="-WP_DEST-"), sg.FolderBrowse()],
    [sg.Text("Replacement Pairs:")],
    [sg.Table(values=[], headings=["Old Text", "New Text"], key="-TABLE-", 
              enable_events=True, auto_size_columns=True, justification='left', 
              num_rows=5, select_mode=sg.TABLE_SELECT_MODE_BROWSE)],
    [sg.Text("Old Text:"), sg.Input(size=(25, 1), key="-OLD_TEXT-"), 
     sg.Text("New Text:"), sg.Input(size=(25, 1), key="-NEW_TEXT-")],
    [sg.Button("Add", key="-ADD-"), sg.Button("Remove", key="-REMOVE-"), 
     sg.Button("Clear All", key="-CLEAR-")],
    [sg.Button("Run Word Processing", key="-RUN_WP-")],
    [sg.ProgressBar(100, orientation='h', size=(40, 20), key='-WP_PROGRESS-')]
]

# Layout for PDF Combine
tab2_layout = [
    [sg.Text("PDF Combine", font=("Any", 16))],
    [sg.Text("DOCX (Word) Origin Folder:"), sg.Input(key="-PC_WORD-"), sg.FolderBrowse()],
    [sg.Text("PDF Invoice Origin Folder:"), sg.Input(key="-PC_PDF-"), sg.FolderBrowse()],
    [sg.Text("Destination Folder:"), sg.Input(key="-PC_DEST-"), sg.FolderBrowse()],
    [sg.Button("Run PDF Combine", key="-RUN_PC-")],
    [sg.ProgressBar(100, orientation='h', size=(40, 20), key='-PC_PROGRESS-')]
]

# Layout for Change Font
tab3_layout = [
    [sg.Text("Change Font in Word Documents", font=("Any", 16))],
    [sg.Text("Origin Folder:"), sg.Input(key="-CF_ORIGIN-"), sg.FolderBrowse()],
    [sg.Text("Destination Folder:"), sg.Input(key="-CF_DEST-"), sg.FolderBrowse()],
    [sg.Button("Run Font Change", key="-RUN_CF-")],
    [sg.ProgressBar(100, orientation='h', size=(40, 20), key='-CF_PROGRESS-')]
]

layout = [
    [sg.TabGroup([[sg.Tab("Word Processing", tab1_layout), 
                  sg.Tab("PDF Combine", tab2_layout), 
                  sg.Tab("Change Font", tab3_layout)]])],
    [sg.Text("Output:")],
    [sg.Multiline(size=(100, 20), key="-OUTPUT-", autoscroll=True, disabled=True)]
]

window = sg.Window("Document Processing GUI", layout, finalize=True)

# Queues for logging and progress updates
log_queue = Queue()
progress_queue = Queue()

# Logging setup
logging.basicConfig(level=logging.INFO, format="%(message)s")

def log(message):
    """Log a message to the GUI and console."""
    logging.info(message)
    log_queue.put(message)

def update_progress(progress_key, value):
    """Update the progress bar."""
    progress_queue.put((progress_key, value))

def validate_folder(path, description):
    """Validate if a folder path exists."""
    if not Path(path).exists():
        log(f"❌ Invalid {description} folder: {path}")
        return False
    return True

def process_word_document(input_path, output_path, replacement_pairs):
    """Process a Word document by replacing text."""
    doc = Document(input_path)
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacement_pairs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
    doc.save(output_path)

def run_word_processing(origin_folder, dest_folder, replacement_pairs):
    """Run Word Processing task."""
    if not replacement_pairs:
        log("❌ Add replacement pairs first!")
        return

    if not validate_folder(origin_folder, "origin") or not validate_folder(dest_folder, "destination"):
        return

    files = list(Path(origin_folder).glob("*.docx"))
    total_files = len(files)
    if total_files == 0:
        log("❌ No Word documents found in the origin folder.")
        return

    for i, file in enumerate(files):
        try:
            output_path = Path(dest_folder) / file.name
            process_word_document(file, output_path, replacement_pairs)
            update_progress("-WP_PROGRESS-", int((i + 1) / total_files * 100))
            log(f"✅ Processed {file.name}")
        except Exception as e:
            log(f"❌ Error processing {file.name}: {str(e)}")

    log("✔ Word Processing completed!")

def run_pdf_combine(values):
    """Run PDF Combine task."""
    word_folder = values["-PC_WORD-"]
    pdf_folder = values["-PC_PDF-"]
    dest_folder = values["-PC_DEST-"]

    if not validate_folder(word_folder, "Word origin") or not validate_folder(pdf_folder, "PDF origin") or not validate_folder(dest_folder, "destination"):
        return

    word_files = list(Path(word_folder).glob("*.docx"))
    pdf_files = list(Path(pdf_folder).glob("*.pdf"))
    total_files = len(word_files)

    if total_files == 0 or len(pdf_files) == 0:
        log("❌ No Word or PDF files found in the specified folders.")
        return

    for i, (word_file, pdf_file) in enumerate(zip(word_files, pdf_files)):
        try:
            output_pdf = Path(dest_folder) / f"{word_file.stem}_combined.pdf"
            convert(word_file, output_pdf)
            merger = PdfMerger()
            merger.append(str(output_pdf))
            merger.append(str(pdf_file))
            merger.write(str(output_pdf))
            merger.close()
            update_progress("-PC_PROGRESS-", int((i + 1) / total_files * 100))
            log(f"✅ Combined {word_file.name} and {pdf_file.name}")
        except Exception as e:
            log(f"❌ Error combining {word_file.name}: {str(e)}")

    log("✔ PDF Combine completed!")

def run_font_change(values):
    """Run Font Change task."""
    origin_folder = values["-CF_ORIGIN-"]
    dest_folder = values["-CF_DEST-"]

    if not validate_folder(origin_folder, "origin") or not validate_folder(dest_folder, "destination"):
        return

    files = list(Path(origin_folder).glob("*.docx"))
    total_files = len(files)
    if total_files == 0:
        log("❌ No Word documents found in the origin folder.")
        return

    for i, file in enumerate(files):
        try:
            doc = Document(file)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)  # Change font size to 12
            output_path = Path(dest_folder) / file.name
            doc.save(output_path)
            update_progress("-CF_PROGRESS-", int((i + 1) / total_files * 100))
            log(f"✅ Updated font in {file.name}")
        except Exception as e:
            log(f"❌ Error updating font in {file.name}: {str(e)}")

    log("✔ Font Change completed!")

# Event loop
while True:
    event, values = window.read(timeout=100)

    # Handle log updates
    while not log_queue.empty():
        window["-OUTPUT-"].update(log_queue.get_nowait() + "\n", append=True)

    # Handle progress updates
    while not progress_queue.empty():
        progress_key, progress_value = progress_queue.get_nowait()
        window[progress_key].update(progress_value)

    # Handle events
    if event == sg.WIN_CLOSED:
        break

    # --- Word Processing Tab Events ---
    elif event == "-ADD-":
        old_text = values["-OLD_TEXT-"].strip()
        new_text = values["-NEW_TEXT-"].strip()
        if old_text and new_text:
            current_table = window["-TABLE-"].get()
            new_pair = [old_text, new_text]
            if new_pair not in current_table:
                current_table.append(new_pair)
                window["-TABLE-"].update(values=current_table)
                window["-OLD_TEXT-"].update("")
                window["-NEW_TEXT-"].update("")
                log(f"➕ Added pair: {old_text} → {new_text}")
            else:
                log("⚠ Pair already exists!")
        else:
            log("❌ Both fields are required!")

    elif event == "-REMOVE-":
        selected_rows = values["-TABLE-"]
        if selected_rows:
            current_table = window["-TABLE-"].get()
            for row in sorted(selected_rows, reverse=True):
                deleted_pair = current_table[row]
                del current_table[row]
                log(f"➖ Removed pair: {deleted_pair[0]} → {deleted_pair[1]}")
            window["-TABLE-"].update(values=current_table)

    elif event == "-CLEAR-":
        window["-TABLE-"].update(values=[])
        log("🧹 Cleared all replacement pairs.")

    # --- Task Execution Events ---
    elif event == "-RUN_WP-":
        # Get data directly from the table element
        replacement_pairs = window["-TABLE-"].get()
        origin_folder = values["-WP_ORIGIN-"]
        dest_folder = values["-WP_DEST-"]
        
        # Validate inputs before starting thread
        if not replacement_pairs:
            log("❌ Add replacement pairs first!")
        elif not origin_folder or not dest_folder:
            log("❌ Origin/Destination folders are required!")
        else:
            threading.Thread(
                target=run_word_processing,
                args=(origin_folder, dest_folder, replacement_pairs),
                daemon=True
            ).start()

    elif event == "-RUN_PC-":
        threading.Thread(target=run_pdf_combine, args=(values,), daemon=True).start()

    elif event == "-RUN_CF-":
        threading.Thread(target=run_font_change, args=(values,), daemon=True).start()

window.close()