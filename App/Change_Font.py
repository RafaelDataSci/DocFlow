#!/usr/bin/env python
"""
Change Font Script

This script reads DOCX files from an origin folder, changes the font of all text to Calibri 12pt
(with the exception of footer text, which is set to Calibri 8pt) while preserving other formatting,
and saves the updated document in a destination folder.

It processes:
    - All paragraphs in the document.
    - All paragraphs within tables.
    - All paragraphs in the footers of each section (with font size 8pt).
"""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)

def change_font_in_docx(file_path: Path, dest_folder: Path) -> None:
    """
    Open a DOCX file, update all text runs to use Calibri 12pt (footer text will be Calibri 8pt),
    and save the modified document to the destination folder.

    Args:
        file_path (Path): The DOCX file to process.
        dest_folder (Path): The folder where the updated DOCX file will be saved.
    """
    try:
        doc = Document(file_path)
        
        # Process paragraphs outside of tables.
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(12)
        
        # Process paragraphs within tables.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(12)
        
        # Process paragraphs in footers (for each section) with font size 8.
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8)
        
        # Ensure the destination folder exists.
        dest_folder.mkdir(parents=True, exist_ok=True)
        new_file_path = dest_folder / file_path.name
        doc.save(new_file_path)
        logging.info(f"Processed file saved as: {new_file_path}")
    
    except Exception as e:
        logging.error(f"Error processing file {file_path}: {e}")

def process_folder(origin_folder: Path, dest_folder: Path) -> None:
    """
    Process all DOCX files in the origin folder.

    Args:
        origin_folder (Path): Folder containing DOCX files.
        dest_folder (Path): Folder where the updated DOCX files will be saved.
    """
    docx_files = list(origin_folder.glob("*.docx"))
    if not docx_files:
        logging.info("No DOCX files found in the origin folder.")
        return

    logging.info(f"Found {len(docx_files)} DOCX file(s) to process.")
    for file in docx_files:
        logging.info(f"Processing {file.name} ...")
        change_font_in_docx(file, dest_folder)
    
    logging.info("All files processed.")

if __name__ == "__main__":
    origin = input("Enter the origin folder path containing DOCX files: ").strip()
    dest = input("Enter the destination folder path for updated DOCX files: ").strip()
    
    origin_folder = Path(origin)
    dest_folder = Path(dest)
    
    if not origin_folder.is_dir():
        logging.error(f"Origin folder does not exist: {origin_folder}")
    else:
        process_folder(origin_folder, dest_folder)
